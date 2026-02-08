# -*- coding: utf-8 -*-

import streamlit as st
from docx import Document
import datetime,time
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_BREAK
from docx.shared import Pt
import re
from num2words import num2words
from io import BytesIO

# -----------------------
# session state variables
# -----------------------

if "infile" not in st.session_state:
    st.session_state.infile = None

# if "totaldownloads" not in st.session_state:
#    st.session_state["totaldownloads"] = 0

infiletext = None
dct_infile = {"costplus":  "costplus.docx", "itemwise": "itemwise.docx", "lumpsum": "lumpsum.docx"}

if "replacements" not in st.session_state:
    st.session_state.replacements = None
    
lov_contract_types = ["Please Select", "CostPlus","ItemWise","LumpSum"]

    
currencies = [ "INR - Indian Rupees", "AED - United Arab Emirates Dirham", "AFN - Afghan Afghani", "ALL - Albanian Lek", "AMD - Armenian Dram",
    "ANG - Netherlands Antillean Guilder", "AOA - Angolan Kwanza", "ARS - Argentine Peso", "AUD - Australian Dollar", "AWG - Aruban Florin",
    "AZN - Azerbaijani Manat", "BAM - Bosnia and Herzegovina Convertible Mark", "BBD - Barbadian Dollar", "BDT - Bangladeshi Taka",
    "BGN - Bulgarian Lev", "BHD - Bahraini Dinar", "BIF - Burundian Franc", "BMD - Bermudian Dollar", "BND - Brunei Dollar", "BOB - Bolivian Boliviano",
    "BRL - Brazilian Real", "BSD - Bahamian Dollar", "BTN - Bhutanese Ngultrum", "BWP - Botswanan Pula", "BYN - Belarusian Ruble", "BZD - Belize Dollar",
    "CAD - Canadian Dollar", "CDF - Congolese Franc", "CHF - Swiss Franc", "CLP - Chilean Peso", "CNY - Chinese Yuan", "COP - Colombian Peso",
    "CRC - Costa Rican Colón", "CUP - Cuban Peso", "CVE - Cape Verdean Escudo", "CZK - Czech Koruna", "DJF - Djiboutian Franc", "DKK - Danish Krone",
    "DOP - Dominican Peso", "DZD - Algerian Dinar",    "EGP - Egyptian Pound","ERN - Eritrean Nakfa","ETB - Ethiopian Birr","EUR - Euro",
    "FJD - Fijian Dollar","FKP - Falkland Islands Pound","GBP - British Pound Sterling","GEL - Georgian Lari","GHS - Ghanaian Cedi",
    "GIP - Gibraltar Pound","GMD - Gambian Dalasi","GNF - Guinean Franc","GTQ - Guatemalan Quetzal","GYD - Guyanaese Dollar","HKD - Hong Kong Dollar",
    "HNL - Honduran Lempira","HRK - Croatian Kuna","HTG - Haitian Gourde","HUF - Hungarian Forint","IDR - Indonesian Rupiah",
    "ILS - Israeli New Shekel","IQD - Iraqi Dinar","IRR - Iranian Rial","ISK - Icelandic Króna","JMD - Jamaican Dollar","JOD - Jordanian Dinar",
    "JPY - Japanese Yen","KES - Kenyan Shilling","KGS - Kyrgyzstani Som","KHR - Cambodian Riel","KMF - Comorian Franc","KPW - North Korean Won",
    "KRW - South Korean Won","KWD - Kuwaiti Dinar","KYD - Cayman Islands Dollar","KZT - Kazakhstani Tenge","LAK - Laotian Kip",
    "LBP - Lebanese Pound","LKR - Sri Lankan Rupee","LRD - Liberian Dollar","LSL - Lesotho Loti","LYD - Libyan Dinar","MAD - Moroccan Dirham",
    "MDL - Moldovan Leu","MGA - Malagasy Ariary","MKD - Macedonian Denar","MMK - Myanmar Kyat","MNT - Mongolian Tugrik","MOP - Macanese Pataca",
    "MRU - Mauritanian Ouguiya","MUR - Mauritian Rupee","MVR - Maldivian Rufiyaa","MWK - Malawian Kwacha","MXN - Mexican Peso","MYR - Malaysian Ringgit",
    "MZN - Mozambican Metical","NAD - Namibian Dollar","NGN - Nigerian Naira","NIO - Nicaraguan Córdoba","NOK - Norwegian Krone","NPR - Nepalese Rupee",
    "NZD - New Zealand Dollar","OMR - Omani Rial","PAB - Panamanian Balboa","PEN - Peruvian Sol","PGK - Papua New Guinean Kina","PHP - Philippine Peso",
    "PKR - Pakistani Rupee","PLN - Polish Zloty","PYG - Paraguayan Guarani","QAR - Qatari Rial","RON - Romanian Leu","RSD - Serbian Dinar",
    "RUB - Russian Ruble","RWF - Rwandan Franc","SAR - Saudi Riyal","SBD - Solomon Islands Dollar","SCR - Seychellois Rupee","SDG - Sudanese Pound",
    "SEK - Swedish Krona","SGD - Singapore Dollar","SLL - Sierra Leonean Leone","SOS - Somali Shilling","SRD - Surinamese Dollar",
    "SSP - South Sudanese Pound",    "STN - São Tomé and Príncipe Dobra","SYP - Syrian Pound","SZL - Swazi Lilangeni","THB - Thai Baht", 
    "TJS - Tajikistani Somoni","TMT - Turkmenistani Manat","TND - Tunisian Dinar","TOP - Tongan Paʻanga","TRY - Turkish Lira",
    "TTD - Trinidad and Tobago Dollar","TWD - New Taiwan Dollar","TZS - Tanzanian Shilling","UAH - Ukrainian Hryvnia","UGX - Ugandan Shilling",
    "USD - United States Dollar","UYU - Uruguayan Peso","UZS - Uzbekistani Som","VES - Venezuelan Bolívar","VND - Vietnamese Dong",
    "VUV - Vanuatu Vatu","WST - Samoan Tala","XAF - Central African CFA Franc","XCD - East Caribbean Dollar","XOF - West African CFA Franc",
    "XPF - CFP Franc","YER - Yemeni Rial","ZAR - South African Rand","ZMW - Zambian Kwacha","ZWL - Zimbabwean Dollar"
]

# Send an email to track count of downloads
def send_email_resend(to_email, subject, text):
    try:
        api_key = os.environ["RESEND_API_KEY"]
        response = requests.post(
            "https://api.resend.com/emails",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json"
            },
            json={
                "from": "onboarding@resend.dev",   # default test sender
                "to": [to_email],
                "subject": subject,
                "text": text
            }
        )
    
        if response.status_code != 200:
            ret = f"ERROR: Email failed ❌ : {response.text}"
        else:
            ret = "SUCCESS: " + response.json()
    
    except Exception as e:
        ret = "EXCEPTION: " + str(e)    

    return(ret)
    
# The output document shd have the same page settings as the input document
def DocSettings(indoc,outdoc):
    
    infile = st.session_state.infile

    indoc = Document(infile)
    text = [p.text for p in indoc.paragraphs]
    # outdoc = Document()

    # global setting for the document style
    src = indoc.sections[0]
    dst = outdoc.sections[0]

    dst.page_width = src.page_width
    dst.page_height = src.page_height

    dst.left_margin = src.left_margin
    dst.right_margin = src.right_margin
    dst.top_margin = src.top_margin
    dst.bottom_margin = src.bottom_margin

    dst.header_distance = src.header_distance
    dst.footer_distance = src.footer_distance
    dst.orientation = src.orientation

    style = outdoc.styles["No Spacing"]
    
    return(outdoc)

# Get the values from within a set of tags
def GetValuesFromTags(txt,list_of_tags):

  replacements = st.session_state.replacements
  
  for tag in list_of_tags:
      if tag in replacements:
        value = replacements[tag]
        txt = txt.replace(tag,value)
  return(txt)

# Format the table data
def AddDataToTable(outdoc,txt):
  totalrows = int(re.findall("<rows=(\\d)+>",txt)[0])

  tabledata = re.findall("<tabledata>(.*?)</tabledata>",txt,re.DOTALL)[0]
  cols = tabledata.split("$COLSEP$")
  totalcols = len(cols)

  tbl = outdoc.add_table(rows=totalrows, cols=totalcols)
  tbl.autofit = True
  tbl.style = "Table Grid"

  for col in range(totalcols):
    entities = cols[col].split("~")
    for entity in entities:
      k,v = entity.split("=")
      k = k.strip(); v = v.strip()

      # key
      tags = re.findall("<(\\w+)>", k) # get all the formatting tags like bold,italic etc.

      CELL = tbl.cell(totalrows-1,col).paragraphs[0]
      if len(tags) > 0:
        k = re.findall("(\\w+)<",k)[0]
        R = CELL.add_run(k)
        R.bold = True
      else:
        R = CELL.add_run(k.strip())
        R.bold = False

      R.font.size = Pt(9) # font size for the Key
      
      # values
      v_value, v_spaces = v.split("^")

      if len(v_value) > 0:
        if "<B>" in v_value:
          v_value = re.findall(">(.*?)<",v_value)[0]

        R = CELL.add_run(" " + v_value)
        R.font.size = Pt(9) # font size for the Value
      
      for _ in range(int(v_spaces) ):
         R = CELL.add_run().add_break(WD_BREAK.LINE)
      
  return(outdoc)

def Spaces():
    st.markdown(""" <hr style = "margin-top: -1tpx; margin-bottom: 0px; } </style> """, unsafe_allow_html=True)
    
def CreateContract():
    Spaces()
    
    c1,c2,c3 = st.columns([0.4,0.2,0.5])
    c1.badge("Select the Contract Type for which you want to generate the Contract",icon=":material/check:")
    contract_type = c2.selectbox(" ", lov_contract_types,label_visibility="collapsed")
    
    
    with st.expander("Contract Conditions and Parties"):
        c1,c2 = st.columns([0.8,0.2])
        # contractname = c1.text_input("Contract Name", max_chars=50,value="Default Contract")
        contractdate = c1.date_input("Contract Date",format="DD/MM/YYYY")
        contractdate = contractdate.strftime("%d/%m/%Y")
        
        c1,c2 = st.columns([0.3,0.7])
        clientname = c1.text_input("Client Name", max_chars=100,value="MyClient")
        clientaddress = c2.text_input("Client Address", max_chars=200,value = "MyClient Address")
        contractorname = c1.text_input("Contractor Name", max_chars=100,value="MyContractor")
        contractoraddress = c2.text_input("Contractor Address", max_chars=200,value="MyContractor Address")
    
    with st.expander("1.Project Details"):
        c1,c2,c3 = st.columns([0.5,0.25,0.25])
        projectdescription = c1.text_input("Project Description", max_chars = 250,value="Default Project")
        projectlocation = c2.text_input("Project Location", max_chars=100,value="Project Location")
        architectengineer = c3.text_input("Architect/Engineer", max_chars=100,value="Architect")
    
    with st.expander("2.Contract Sum and Payments"):
        c1,c2,c3,c4 = st.columns(4)
        currency = c1.selectbox("Currency", currencies)
        contractvalue = c2.number_input("Contract Value", min_value=1)
        markup = c3.number_input("Markup %", min_value=1)
        
        c1,c2,c3,c4 = st.columns(4)
        advanceamountfromclient = c1.number_input("Advance Amount from Client")
        paymentprogress = c2.selectbox("Cost Statement submission (Contractor)",["Monthly","Biweekly","Weekly"])
        finalbillfromcontractor = c3.number_input("Final Bill from Contractor (Days)", min_value=1,value=30)
        paymentrelease = c4.number_input("Payment Release from Client (Days)", min_value=1,value=30)
        
    with st.expander("3.Commencement,Progress,Completion"):
        c1,c2,c3 = st.columns(3)
        commencementdate = c1.date_input("Commencement Date", format="DD/MM/YYYY")
        commencementdate = commencementdate.strftime("%d/%m/%Y")
        completionperiod = c2.number_input("Completion Period (Days)", min_value=1,value=30)
        termination = c3.number_input("Termination (Days)", min_value=1)
    
    with st.expander("4.Variations and Cost Adjustments"):
        c1,c2,c3 = st.columns([0.3,0.3,0.4])
        changerecord = c1.number_input("Change Record within (Days)", min_value=1)
        nextschedulefinal = c2.selectbox("Cost Adjustment Claim",["Next Schedule","Final"])
        
    with st.expander("5.Performance Security and Defects Liability"):
        c1,c2,c3,c4 = st.columns([0.3,0.3,0.3,0.1])
        holdamount = c1.number_input("Hold Amount", min_value=0)
        retentionperiod = c2.number_input("Retention Period", min_value=0)
        retentionlimit = c3.number_input("Retention Limit (Days)", min_value=0)
        
    with st.expander("6.Suspension and Termination"):
        c1,c2,c3 = st.columns(3)
        abandon = c1.number_input("Abandon Work [Days]",min_value=1)
        clientduepayment = c2.number_input("Client Payment Due (Days)", min_value=1,value=30)
        finalsettlement = c3.number_input("Final Settlement (Days)", min_value=1,value=45)
        
    with st.expander("8.Dispute Resolution"):
        c1,c2,c3 = st.columns(3)
        dispute = c1.number_input("Unresolved Period (Days)",min_value=1,value=7)
        arbitrationlocation = c2.text_input("Arbitration Location", max_chars=100,value="Location")
        language = c3.selectbox("Proceedings Language", ["English","Tamil"])
        
    with st.expander("10.Governing Law"):
        c1,c2 = st.columns(2)
        country = c1.text_input("Governing Country", max_chars=100,value="India")
        jurisdictioncity = c2.text_input("Jurisdiction City", max_chars=100,value="City")
        
    btn_create = st.button(":page_with_curl:",help="Prepare Contract")

    if btn_create:
        if contract_type == "Please Select":
            st.error("Select a valid Contract Type before creating the Contract")
            return
        
        replacements = { 
            # "[contractname]": "<B>"+contractname+"</B>",
            "[contractdate]": "<B>"+contractdate+"</B>", 
            "[clientname]": "<B>"+clientname+"</B>", 
            "[clientaddress]": "<B>"+clientaddress+"</B>", 
            "[contractorname]": "<B>"+contractorname+"</B>", 
            "[contractoraddress]": "<B>"+contractoraddress+"</B>", 
            "[projectdescription]": "<B>"+projectdescription+"</B>", 
            "[projectlocation]": "<B>"+projectlocation+"</B>", 
            "[architectengineer]": "<B>"+architectengineer+"</B>", 
            "[currency]": "<B>"+currency+"</B>",
            "[contractvalue]": "<B>"+str(contractvalue)+"</B>",
            "[contractvaluewords]": "<B>"+num2words(contractvalue)+"</B>",
            "[markup]": "<B>"+str(markup)+"</B>", 
            "[advanceamountfromclient]": "<B>"+str(advanceamountfromclient)+"</B>", 
            "[paymentprogress]": "<B>"+paymentprogress+"</B>", 
            "[finalbillfromcontractor]": "<B>"+str(finalbillfromcontractor)+"</B>",
            "[paymentrelease]": "<B>"+str(paymentrelease)+"</B>",
            "[commencementdate]": "<B>"+commencementdate+"</B>",
            "[completionperiod]": "<B>"+str(completionperiod)+"</B>",
            "[termination]": "<B>"+str(termination)+"</B>",
            "[changerecord]": "<B>"+str(changerecord)+"</B>", 
            "[nextschedulefinal]": "<B>"+nextschedulefinal+"</B>", 
            "[holdamount]": "<B>"+str(holdamount)+"</B>",
            "[retentionperiod]": "<B>"+str(retentionperiod)+"</B>",
            "[retentionlimit]": "<B>"+str(retentionlimit)+"</B>",
            "[abandon]": "<B>"+str(abandon)+"</B>", 
            "[clientduepayment]": "<B>"+str(clientduepayment)+"</B>", 
            "[finalsettlement]": "<B>"+str(finalsettlement)+"</B>", 
            "[dispute]": "<B>"+str(dispute)+"</B>",
            "[arbitrationlocation]": "<B>"+arbitrationlocation+"</B>", 
            "[language]": "<B>"+language+"</B>", 
            "[country]": "<B>"+country+"</B>", 
            "[jurisdictioncity]": "<B>"+jurisdictioncity+"</B>", 
            }
            
        st.session_state.replacements = replacements
        
        infile = dct_infile[contract_type.strip().lower() ]
        st.session_state.infile = infile

        with st.status("Initializing Contract Document preparation. Please wait ... ", expanded=True) as status:
            time.sleep(1)
        
            # file name will start with the current date and time
            today = datetime.datetime.now()
            today = datetime.datetime.strftime(today,"%d%m%Y") + str(int(time.time()))
            
            indoc = Document(infile)
            text = [p.text for p in indoc.paragraphs]
            
            outfile = f"{today}_{contract_type.strip().lower()}_{clientname}.docx"
            
            status.update(label="Preparing Contract ...", state="running")
            time.sleep(1)

            outdoc = Document()
            outdoc = DocSettings(indoc,outdoc)

            # global setting for the document style
            # p = outdoc.add_paragraph(style = "No Spacing")
            
            # Main Code
            for i in range(len(text)):
                P = outdoc.add_paragraph(style = "No Spacing") # added 9/1/2026
                
                txt = text[i]

                if txt == "[EOD]":
                   pass
                elif txt == "[EOP]": # or len(txt.strip())<= 0:
                   # outdoc.add_page_break() # original code. commeted 9/1/2026
                   
                   # added 9/1/2026
                   p = outdoc.add_paragraph()
                   # P.add_run().add_break(WD_BREAK.PAGE)
                   if outdoc.paragraphs:
                        outdoc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
                   else:
                        outdoc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                   # addition 9/1/2026 ends here
                   # continue
                   
                # if the input data is a table data
                elif "[TABLE]" in txt:
                    list_of_tags = re.findall(r"\[\w+\]",txt)
                    txt = GetValuesFromTags(txt,list_of_tags)
                    outdoc = AddDataToTable(outdoc,txt)
                # normal text / data
                else:
                    # extract all the tags that needs formatting. eg: bold, italics, alignments etc.
                    formatting_tags = re.findall(r'<([^>]+)>', txt)
                    if len(formatting_tags) > 0:
                        formatting_tags = formatting_tags[0]

                        # extract the text that needs formatting
                        txt = re.findall("([^<]+)<",txt)[0]
                        # P = outdoc.add_paragraph(style = "No Spacing") # comment by CRS 9/1/2026
                        R = P.add_run(txt)
                        P.runs[0].font.size = Pt(9)
                        
                        # check the different types of tags that can be part of the formatting
                        # based on that, apply that setting on the text

                        if 'center' in formatting_tags:
                            P.alignment = WD_ALIGN_PARAGRAPH.CENTER

                        if 'left' in formatting_tags:
                            P.alignment = WD_ALIGN_PARAGRAPH.LEFT

                        if 'bold' in formatting_tags:
                            R.bold = True
                        
                        if 'bullet' in formatting_tags:
                            P.style = "List Bullet"

                        if 'italic' in formatting_tags:
                            R.italic = True
                            
                        if 'underline' in formatting_tags:
                            R.underline = True

                        # set the font size
                        if "font" in formatting_tags:
                            fs = int(re.findall("font:(\\d+),",formatting_tags)[0])
                            P.runs[0].font.size = Pt(fs)

                        if "spaces" in formatting_tags:
                            sp = int(re.findall("spaces:(\\d+),",formatting_tags)[0])
                            for _ in range(sp):
                                # r_space = P.add_run("\n")
                                r_space = P.add_run().add_break(WD_BREAK.LINE)
                    else:
                        # P = outdoc.add_paragraph(style = "No Spacing")
                        
                        list_of_tags = re.findall(r"\[\w+\]",txt)
                        if len(list_of_tags) > 0:
                            txt = GetValuesFromTags(txt,list_of_tags)

                        start_bold = txt.split("<B>")
                        for start in start_bold:
                            if "</B>" not in start:
                                r = P.add_run(start)
                                r.font.size = Pt(9)
                            else:
                                splits = start.split("</B>")
                                for i in range(len(splits)):
                                    R = P.add_run(splits[i])
                                    R.font.size = Pt(9)
                                    if i == 0:
                                        R.bold = True
                                    else:
                                        R.bold = False

            status.update(label="Contract created successfully!!!", state="complete")
            time.sleep(1)

            # dont save a physical copy of the file. Store in the memory
            buffer = BytesIO()
            outdoc.save(buffer)
            buffer.seek(0)
            
            # download filename
            download_file = f"{today}_{contract_type.strip().lower()}_{clientname}.docx"
            
            btn_download = st.download_button(  label="Download Contract", 
                                                icon=":material/download:", 
                                                data=buffer,
                                                file_name=download_file,
                                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", )

            
            st.session_state["totaldownloads"]+=1

            to_email = "copactbeta@gmail.com"
            subject = "Contract: " + contract_type
            text = "Contracted downloaded"
            email_status = send_email_resend(to_email, subject, text)

            st.write("Email status = ", email_status)

def main():
    CreateContract()









